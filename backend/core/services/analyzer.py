"""
Servicio de análisis de RFPs usando Gemini via Google AI API.
Trabaja con archivos locales extrayendo texto primero.
Soporta grounding para obtener tarifas de mercado actuales.
"""
import logging
from datetime import datetime
from pathlib import Path
from typing import Any, Literal

from pypdf import PdfReader
from docx import Document
import io

from core.gcp.gemini_client import get_gemini_client
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from models.certification import Certification

logger = logging.getLogger(__name__)

# Cargar prompts
PROMPTS_DIR = Path(__file__).parent.parent.parent / "prompts"


def load_prompt(name: str) -> str:
    """Carga un prompt desde archivo."""
    prompt_path = PROMPTS_DIR / f"{name}.txt"
    if prompt_path.exists():
        return prompt_path.read_text(encoding="utf-8")
    
    # Prompt por defecto si no existe el archivo
    if name == "rfp_analysis":
        return DEFAULT_ANALYSIS_PROMPT
    elif name == "question_generation":
        return DEFAULT_QUESTIONS_PROMPT
    
    raise FileNotFoundError(f"Prompt not found: {prompt_path}")


DEFAULT_ANALYSIS_PROMPT = """
Eres un experto analista de RFPs (Request for Proposals) para TIVIT, una empresa líder en tecnología y servicios digitales en Latinoamérica.

Analiza el siguiente documento RFP y extrae la información estructurada en formato JSON con el siguiente schema:

{
    "title": "Título oficial del proyecto o licitación",
    "client_name": "Nombre COMPLETO del cliente/empresa",
    "client_acronym": "Siglas o abreviatura del cliente (ej: BCP, INDAP, BBVA)",
    "country": "País del cliente",
    "category": "Categoría del proyecto (infraestructura, desarrollo, cloud, seguridad, etc.)",
    "summary": "Resumen ejecutivo del RFP en máximo 3 párrafos",
    "scope": {
        "description": "Descripción detallada del alcance",
        "deliverables": ["Lista de entregables esperados"],
        "exclusions": ["Lo que NO está incluido"]
    },
    "requirements": {
        "technical": ["Requisitos técnicos"],
        "functional": ["Requisitos funcionales"],
        "certifications": ["Certificaciones requeridas"]
    },
    "budget": {
        "amount_min": null,
        "amount_max": null,
        "currency": "USD",
        "notes": "Notas sobre presupuesto"
    },
    "proposal_deadline": "YYYY-MM-DD",
    "questions_deadline": "YYYY-MM-DD",
    "project_duration": "Duración estimada del proyecto",
    "evaluation_criteria": [
        {"criterion": "nombre", "weight": 0.0}
    ],
    "risks": [
        {"risk": "descripción", "severity": "high/medium/low", "mitigation": "posible mitigación"}
    ],
    "opportunities": ["Oportunidades identificadas para TIVIT"],
    "confidence_score": 0.0,
    "recommendation": "GO o NO_GO con justificación breve"
}

Sé preciso y extrae toda la información relevante. Si algún campo no está disponible, usa null.
"""


DEFAULT_QUESTIONS_PROMPT = """
Eres un experto en desarrollo de negocios de TIVIT que prepara preguntas estratégicas para enviar al cliente antes de presentar una propuesta.

Basándote en el análisis del RFP, genera entre 10 y 15 preguntas inteligentes que:
1. Clarifiquen ambigüedades en los requisitos
2. Identifiquen oportunidades de valor agregado
3. Ayuden a preparar una propuesta más competitiva
4. Demuestren expertise y profesionalismo

Responde con un JSON array de objetos:
[
    {
        "question": "La pregunta completa",
        "category": "técnica|comercial|alcance|timeline|legal",
        "priority": "alta|media|baja",
        "context": "Por qué es importante esta pregunta",
        "why_important": "Cómo impacta en la propuesta"
    }
]

Ordena las preguntas por prioridad (alta primero).
"""


class RFPAnalyzerService:
    """Servicio para analizar RFPs con Gemini."""
    
    def __init__(self):
        self._gemini = None
        try:
            self.analysis_prompt = load_prompt("rfp_analysis")
        except FileNotFoundError:
            self.analysis_prompt = DEFAULT_ANALYSIS_PROMPT
        
        try:
            self.questions_prompt = load_prompt("question_generation")
        except FileNotFoundError:
            self.questions_prompt = DEFAULT_QUESTIONS_PROMPT
        
        try:
            self.certification_prompt = load_prompt("certification_analysis")
        except FileNotFoundError:
            self.certification_prompt = "Analiza este certificado y extrae: name, issuer, description, issue_date, expiry_date, scope en JSON."

        try:
            self.chapter_prompt = load_prompt("chapter_analysis")
        except FileNotFoundError:
            logger.warning("Chapter prompt not found, using fallback")
            self.chapter_prompt = "Analiza este capítulo de propuesta técnica y extrae: name (título del capítulo), description (breve resumen del contenido) en JSON."

    async def analyze_certification_content(self, content: bytes, filename: str) -> dict[str, Any]:
        """
        Analiza un documento de certificación.
        """
        logger.info(f"Analyzing certification: {filename}")
        document_text = self.extract_text(content, filename)
        
        if not document_text.strip():
            return {"name": filename, "description": "No text extracted"}
            
        result = await self.gemini.analyze_document(
            document_content=document_text,
            prompt=self.certification_prompt,
            analysis_mode="fast"
        )
        return result

    async def analyze_chapter_content(self, content: bytes, filename: str) -> dict[str, Any]:
        """
        Analiza un documento de capítulo.
        """
        logger.info(f"Analyzing chapter: {filename}")
        document_text = self.extract_text(content, filename)
        
        if not document_text.strip():
            return {"name": filename, "description": "No text extracted"}
            
        result = await self.gemini.analyze_document(
            document_content=document_text,
            prompt=self.chapter_prompt,
            analysis_mode="fast"
        )
        return result
    
    @property
    def gemini(self):
        """Lazy loading del cliente Gemini."""
        if self._gemini is None:
            self._gemini = get_gemini_client()
        return self._gemini
    
    def extract_text_from_pdf(self, content: bytes) -> str:
        """Extrae texto de un PDF."""
        try:
            reader = PdfReader(io.BytesIO(content))
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise
    
    def extract_text_from_docx(self, content: bytes) -> str:
        """Extrae texto de un DOCX."""
        try:
            doc = Document(io.BytesIO(content))
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # También extraer de tablas
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise
    
    def extract_text(self, content: bytes, filename: str) -> str:
        """Extrae texto de un archivo según su extensión."""
        filename_lower = filename.lower()
        
        if filename_lower.endswith(".pdf"):
            return self.extract_text_from_pdf(content)
        elif filename_lower.endswith(".docx"):
            return self.extract_text_from_docx(content)
        else:
            # Asumir texto plano
            return content.decode("utf-8", errors="ignore")
    
    async def analyze_rfp_from_content(
        self, 
        content: bytes, 
        filename: str,
        analysis_mode: Literal["fast", "balanced", "deep"] = "balanced",
        use_grounding: bool = True,
        db: AsyncSession | None = None,
    ) -> dict[str, Any]:
        """
        Analiza un RFP desde su contenido en bytes.
        
        Args:
            content: Contenido del archivo en bytes
            filename: Nombre del archivo (para determinar tipo)
            analysis_mode: Modo de análisis (fast/balanced/deep)
            use_grounding: Si True, usa Google Search para tarifas de mercado
            db: Sesión de base de datos para obtener certificaciones
            
        Returns:
            Datos extraídos del RFP incluyendo team_estimation y cost_estimation
        """
        logger.info(f"Starting RFP analysis for: {filename}")
        logger.info(f"Analysis mode: {analysis_mode}, Grounding: {use_grounding}")
        
        # Extraer texto del documento
        document_text = self.extract_text(content, filename)
        
        if not document_text.strip():
            logger.error("No text extracted from document")
            return {"error": "No se pudo extraer texto del documento"}
        
        logger.info(f"Extracted {len(document_text)} characters from document")
        
        # Preparar prompt con certificaciones si hay DB
        prompt_to_use = self.analysis_prompt
        
        if db:
            try:
                # Obtener certificaciones activas
                result = await db.execute(select(Certification).where(Certification.is_active == True))
                certs = result.scalars().all()
                if certs:
                    cert_list_str = "\n".join([f"- {c.name} (ID: {c.id}): {c.description[:100]}..." for c in certs])
                    prompt_to_use = prompt_to_use.replace("{{available_certifications}}", cert_list_str)
                    logger.info(f"Injected {len(certs)} certifications into prompt")
                else:
                    prompt_to_use = prompt_to_use.replace("{{available_certifications}}", "No hay certificaciones disponibles.")
            except Exception as e:
                logger.error(f"Error fetching certifications for prompt: {e}")
                prompt_to_use = prompt_to_use.replace("{{available_certifications}}", "Error al recuperar certificaciones.")
        else:
            prompt_to_use = prompt_to_use.replace("{{available_certifications}}", "No disponible (sin conexión a DB).")

        # Analizar con Gemini - usar grounding si está habilitado
        if use_grounding:
            logger.info("Using Gemini with Google Search Grounding for market rates")
            result = await self.gemini.analyze_with_grounding(
                document_content=document_text,
                prompt=prompt_to_use,
                temperature=0.1,
                max_output_tokens=16384,
            )
        else:
            result = await self.gemini.analyze_document(
                document_content=document_text,
                prompt=prompt_to_use,
                analysis_mode=analysis_mode,
            )
        
        logger.info(f"RFP analysis completed: {result}")
        return result
    
    async def analyze_rfp(self, gcs_uri: str, use_grounding: bool = True, db: AsyncSession | None = None) -> dict[str, Any]:
        """
        Analiza un RFP desde GCS o local.
        
        Args:
            gcs_uri: URI del archivo en Cloud Storage o local://
            use_grounding: Si True, usa Google Search para tarifas de mercado
            db: Sesión de DB opcional
            
        Returns:
            Datos extraídos del RFP
        """
        logger.info(f"Starting RFP analysis: {gcs_uri}")
        
        # Si es un URI local, leer el archivo
        if gcs_uri.startswith("local://"):
            from core.storage import get_storage_service
            storage = get_storage_service()
            content = storage.download_file(gcs_uri)
            filename = gcs_uri.split("/")[-1]
            return await self.analyze_rfp_from_content(
                content, 
                filename, 
                use_grounding=use_grounding,
                db=db
            )
        
        # Para GCS, descargar y analizar
        from core.storage import get_storage_service
        storage = get_storage_service()
        content = storage.download_file(gcs_uri)
        filename = gcs_uri.split("/")[-1]
        return await self.analyze_rfp_from_content(
            content, 
            filename, 
            use_grounding=use_grounding,
            db=db
        )
    
    async def generate_questions(self, rfp_data: dict[str, Any]) -> list[dict[str, Any]]:
        """
        Genera preguntas basadas en el análisis del RFP.
        
        Args:
            rfp_data: Datos extraídos del análisis
            
        Returns:
            Lista de preguntas generadas
        """
        logger.info("Generating questions for RFP")
        
        questions = await self.gemini.generate_questions(
            rfp_data=rfp_data,
            prompt=self.questions_prompt,
            temperature=0.3,
        )
        
        logger.info(f"Generated {len(questions)} questions")
        return questions
    
    def extract_indexed_fields(self, extracted_data: dict[str, Any]) -> dict[str, Any]:
        """
        Extrae campos para indexación en la base de datos.
        
        Args:
            extracted_data: Datos completos extraídos
            
        Returns:
            Campos para actualizar en el modelo
        """
        # Manejar el campo budget de manera robusta
        budget_data = extracted_data.get("budget")
        if isinstance(budget_data, dict):
            budget = budget_data
        elif isinstance(budget_data, list) and len(budget_data) > 0 and isinstance(budget_data[0], dict):
            # Si es una lista, tomar el primer elemento
            budget = budget_data[0]
        else:
            # Por defecto, diccionario vacío
            budget = {}
        
        # Parsear fechas
        proposal_deadline = None
        questions_deadline = None
        
        if extracted_data.get("proposal_deadline"):
            try:
                proposal_deadline = datetime.strptime(
                    extracted_data["proposal_deadline"], "%Y-%m-%d"
                ).date()
            except ValueError:
                pass
        
        if extracted_data.get("questions_deadline"):
            try:
                questions_deadline = datetime.strptime(
                    extracted_data["questions_deadline"], "%Y-%m-%d"
                ).date()
            except ValueError:
                pass
        
        return {
            "title": extracted_data.get("title"),
            "client_name": extracted_data.get("client_name"),
            "client_acronym": extracted_data.get("client_acronym"),
            "country": extracted_data.get("country"),
            "category": extracted_data.get("category"),
            "summary": extracted_data.get("summary"),
            "budget_min": budget.get("amount_min") if isinstance(budget, dict) else None,
            "budget_max": budget.get("amount_max") if isinstance(budget, dict) else None,
            "currency": budget.get("currency", "USD") if isinstance(budget, dict) else "USD",
            "proposal_deadline": proposal_deadline,
            "questions_deadline": questions_deadline,
            "project_duration": extracted_data.get("project_duration"),
            "confidence_score": extracted_data.get("confidence_score"),
            "recommendation": extracted_data.get("recommendation"),
        }


    async def analyze_experience_relevance(
        self, 
        rfp_summary: str,
        experiences: list[dict[str, Any]]
    ) -> list[dict[str, Any]]:
        """
        Analiza la relevancia de las experiencias para un RFP.
        Ref: Step Id: 2289
        """
        if not experiences:
            return []

        # Prepare prompt
        experiences_text = "\n".join([
            f"ID: {exp['id']}\nCliente: {exp.get('propietario_servicio')}\nDescripción: {exp.get('descripcion_servicio')}\nMonto: {exp.get('monto_final')}\n---"
            for exp in experiences
        ])

        prompt = f"""
        Actúa como un experto en preventa de servicios TI.
        
        RESUMEN DEL RFP:
        {rfp_summary}

        EXPERIENCIAS PREVIAS (Portafolio):
        {experiences_text}

        TAREA:
        Clasifica CADA experiencia según su relevancia para este RFP.
        Genera un JSON list con objetos: {{ "experience_id": "UUID", "score": 0.0-1.0, "reason": "Breve justificación (max 15 palabras)" }}
        
        CRITERIOS:
        - Score > 0.8: Match directo (Misma industria, tecnología similar, o caso de uso idéntico).
        - Score > 0.5: Match parcial (Tecnología similar o industria similar).
        - Score < 0.5: Poca relevancia.
        - Se estricto. No alucines IDs. Usa SOLO los IDs provistos.
        """

        try:
            # Use the new helper method directly
            recommendations = await self.gemini.generate_json(prompt)
            
            # Ensure it's a list
            if isinstance(recommendations, dict):
                recommendations = [recommendations]
            elif not isinstance(recommendations, list):
                logger.warning(f"AI returned unexpected format: {type(recommendations)}")
                return []
            
            # Normalize scores to 0-100 if they came as 0-1
            # But the requirement implies returning list of objects matching schema
            return recommendations

        except Exception as e:
            logger.error(f"Error in analyze_experience_relevance: {e}")
            return []

    async def analyze_chapter_relevance(
        self, 
        rfp_summary: str,
        chapters: list[dict[str, Any]]
    ) -> list[dict[str, Any]]:
        """
        Analiza la relevancia de los capítulos para un RFP.
        """
        if not chapters:
            return []

        # Prepare prompt
        chapters_text = "\n".join([
            f"ID: {chap['id']}\nNombre: {chap['name']}\nDescripción: {chap.get('description', '')}\n---"
            for chap in chapters
        ])

        prompt = f"""
        Actúa como un experto en redacción de propuestas técnicas.
        
        RESUMEN DEL RFP:
        {rfp_summary}

        CAPÍTULOS DISPONIBLES (Biblioteca de Contenidos):
        {chapters_text}

        TAREA:
        Evalúa qué capítulos son útiles para responder a este RFP específico.
        Genera un JSON list con objetos: {{ "chapter_id": "UUID", "score": 0.0-1.0, "reason": "Breve justificación (max 10 palabras)" }}
        
        CRITERIOS:
        - Score > 0.8: Capítulo Esencial (ej: Metodología, Equipo, si el RFP lo pide explícitamente).
        - Score > 0.5: Capítulo Útil (Complementario).
        - Score < 0.5: Irrelevante.
        - Se estricto. Usa SOLO los IDs provistos.
        """

        try:
            recommendations = await self.gemini.generate_json(prompt)
            
            if isinstance(recommendations, dict):
                recommendations = [recommendations]
            elif not isinstance(recommendations, list):
                return []
            
            return recommendations

        except Exception as e:
            logger.error(f"Error in analyze_chapter_relevance: {e}")
            return []
            
# Singleton instance
_analyzer_service: RFPAnalyzerService | None = None


def get_analyzer_service() -> RFPAnalyzerService:
    """Get or create analyzer service singleton."""
    global _analyzer_service
    if _analyzer_service is None:
        _analyzer_service = RFPAnalyzerService()
    return _analyzer_service
